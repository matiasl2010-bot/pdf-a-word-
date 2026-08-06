import re

from docx import Document


def convertir(markdown_text: str, output_path: str) -> None:
    doc = Document()
    lineas = markdown_text.split("\n")
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        if linea.strip().startswith("|"):
            i = _agregar_tabla(doc, lineas, i)
            continue
        _agregar_linea(doc, linea)
        i += 1
    doc.save(output_path)


def _agregar_linea(doc, linea: str) -> None:
    if linea.startswith("## "):
        doc.add_heading(linea[3:].strip(), level=2)
    elif linea.startswith("# "):
        doc.add_heading(linea[2:].strip(), level=1)
    elif linea.startswith("- ") or linea.startswith("* "):
        p = doc.add_paragraph(style="List Bullet")
        _agregar_runs_con_formato(p, linea[2:].strip())
    elif linea.strip():
        p = doc.add_paragraph()
        _agregar_runs_con_formato(p, linea)


def _agregar_runs_con_formato(paragraph, texto: str) -> None:
    partes = re.split(r"(\*\*.+?\*\*)", texto)
    for parte in partes:
        if not parte:
            continue
        if parte.startswith("**") and parte.endswith("**"):
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
        else:
            paragraph.add_run(parte)


def _es_separador_tabla(linea: str) -> bool:
    return bool(re.fullmatch(r"\|[\s\-:|]+\|?", linea.strip()))


def _agregar_tabla(doc, lineas: list, inicio: int) -> int:
    filas_texto = []
    i = inicio
    while i < len(lineas) and lineas[i].strip().startswith("|"):
        if not _es_separador_tabla(lineas[i]):
            celdas = [c.strip() for c in lineas[i].strip().strip("|").split("|")]
            filas_texto.append(celdas)
        i += 1

    if not filas_texto:
        return i

    n_cols = len(filas_texto[0])
    tabla = doc.add_table(rows=0, cols=n_cols)
    tabla.style = "Table Grid"
    for fila in filas_texto:
        celdas_fila = tabla.add_row().cells
        for j in range(n_cols):
            celdas_fila[j].text = fila[j] if j < len(fila) else ""

    return i
