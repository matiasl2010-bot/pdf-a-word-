import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
import markdown_to_docx

MARKDOWN_DE_PRUEBA = """# Titulo principal
Un parrafo con **texto en negrita** normal.
## Subtitulo
- primer item
- segundo item

| Componente | Valor |
| --- | --- |
| Resistencia | 220 ohm |
| Capacitor | 10 uF |
"""


def test_convertir_genera_docx_con_titulos_y_parrafos(tmp_path):
    salida = tmp_path / "salida.docx"

    markdown_to_docx.convertir(MARKDOWN_DE_PRUEBA, str(salida))

    assert salida.exists()
    doc = Document(str(salida))
    estilos = [p.style.name for p in doc.paragraphs if p.text.strip()]
    textos = [p.text for p in doc.paragraphs]
    assert "Heading 1" in estilos
    assert "Heading 2" in estilos
    assert any("Titulo principal" in t for t in textos)
    assert any("Subtitulo" in t for t in textos)


def test_convertir_preserva_negrita(tmp_path):
    salida = tmp_path / "salida.docx"
    markdown_to_docx.convertir(MARKDOWN_DE_PRUEBA, str(salida))

    doc = Document(str(salida))
    parrafo_negrita = next(p for p in doc.paragraphs if "texto en negrita" in p.text)
    runs_bold = [r for r in parrafo_negrita.runs if r.bold and "negrita" in r.text]
    assert len(runs_bold) == 1


def test_convertir_genera_lista(tmp_path):
    salida = tmp_path / "salida.docx"
    markdown_to_docx.convertir(MARKDOWN_DE_PRUEBA, str(salida))

    doc = Document(str(salida))
    items = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert "primer item" in items
    assert "segundo item" in items


def test_convertir_genera_tabla(tmp_path):
    salida = tmp_path / "salida.docx"
    markdown_to_docx.convertir(MARKDOWN_DE_PRUEBA, str(salida))

    doc = Document(str(salida))
    assert len(doc.tables) == 1
    tabla = doc.tables[0]
    assert tabla.cell(0, 0).text == "Componente"
    assert tabla.cell(1, 0).text == "Resistencia"
    assert tabla.cell(1, 1).text == "220 ohm"
